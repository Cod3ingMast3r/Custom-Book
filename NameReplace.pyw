import os
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
import pandas as pd
from docx import Document
import subprocess
import win32com.client
import time

class BookApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Book App")
        self.root.geometry("600x300")

        self.doc = Document()
        self.new_doc = Document()
        self.pdf_export_path = ""

        main_frame = ttk.Frame(root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Dynamic row numbers
        book_row = 0
        name_file_row = book_row + 1
        empty_row_1 = name_file_row + 1
        export_file_row = empty_row_1 + 1
        preview_row = export_file_row + 2

        # Book Selection
        self.book_label = ttk.Label(main_frame, text="Select a Book:")
        self.book_label.grid(row=book_row, column=0, sticky=tk.W)

        self.book_folders = [f for f in os.listdir("Books") if os.path.isdir(os.path.join("Books", f))]
        self.book_var = tk.StringVar()
        self.book_var.set("Select a folder")

        self.book_dropdown = ttk.Combobox(main_frame, textvariable=self.book_var, values=self.book_folders)
        self.book_dropdown.grid(row=book_row, column=1, sticky=(tk.W, tk.E))
        self.book_dropdown.bind("<<ComboboxSelected>>", self.load_book)

        or_label = ttk.Label(main_frame, text="or", justify=tk.CENTER, anchor=tk.CENTER)
        or_label.grid(row=book_row, column=2, sticky=(tk.W, tk.E))

        self.select_book_file_button = ttk.Button(main_frame, text="Select Book File", command=self.select_book_file)
        self.select_book_file_button.grid(row=book_row, column=3, sticky=tk.W)

        # Name File Selection
        self.name_file_label = ttk.Label(main_frame, text="Select a Name File:")
        self.name_file_label.grid(row=name_file_row, column=0, sticky=tk.W)

        self.new_name_files = [f.replace(".xlsx", "") for f in os.listdir("New Names") if ".xlsx" in f]
        self.new_name_file_var = tk.StringVar()
        self.new_name_file_var.set("Select a Name File")

        self.name_file_dropdown = ttk.Combobox(main_frame, textvariable=self.new_name_file_var, values=self.new_name_files)
        self.name_file_dropdown.grid(row=name_file_row, column=1, sticky=(tk.W, tk.E))
        self.name_file_dropdown.bind("<<ComboboxSelected>>", self.load_excel)

        or_label1 = ttk.Label(main_frame, text="or", justify=tk.CENTER, anchor=tk.CENTER)
        or_label1.grid(row=name_file_row, column=2, sticky=(tk.W, tk.E))

        self.select_names_file_button = ttk.Button(main_frame, text="Select Name File", command=self.select_names_file)
        self.select_names_file_button.grid(row=name_file_row, column=3, sticky=tk.W)

        # Empty Row
        empty_label = ttk.Label(main_frame, text="")
        empty_label.grid(row=empty_row_1, column=0, sticky=tk.W)

        # File Name Entry
        self.label3 = ttk.Label(main_frame, text="Export File Name:")
        self.label3.grid(row=export_file_row, column=0, sticky=tk.W)
        self.file_name_var = tk.StringVar()
        self.file_name_entry = ttk.Entry(main_frame, textvariable=self.file_name_var)
        self.file_name_entry.grid(row=export_file_row, column=1, sticky=(tk.W, tk.E))

        self.export_button = ttk.Button(main_frame, text="Export", command=self.export)
        self.export_button.grid(row=export_file_row, column=2, sticky=tk.W)

        # Error Label
        self.error_label = ttk.Label(main_frame, text="", foreground="red")
        self.error_label.grid(row=export_file_row+1, column=0, columnspan=4, sticky=(tk.W, tk.E))

        # Preview Button
        self.preview_button = ttk.Button(main_frame, text="Preview", command=self.preview)
        self.preview_button.grid(row=preview_row, column=0, sticky=tk.W)
        
    # Load the selected book
    def load_book(self, event):
        folder = self.book_var.get()
        doc_path = os.path.join("Books", folder, "book.docx")
        self.doc = Document(doc_path)

    # Load the selected Excel file
    def load_excel(self, event):
        file = self.new_name_file_var.get()
        file_path = os.path.join("New Names", file + ".xlsx")
        self.df = pd.read_excel(file_path)

    # Select a book file manually
    def select_book_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if file_path:
            self.doc = Document(file_path)

    # Select a name file manually
    def select_names_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            self.df = pd.read_excel(file_path)

    # Replace text in the document
    def replace_text_in_document(self):
        self.new_doc = Document()
        for paragraph in self.doc.paragraphs:
            new_paragraph = self.new_doc.add_paragraph()
            for run in paragraph.runs:
                new_run = new_paragraph.add_run()
                new_run.text = run.text
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                for index, row in self.df.iterrows():
                    if row['New'] != "":
                        new_run.text = new_run.text.replace(row['Original'], row['New'])

    # Export the document
    def export(self):
        file_name = self.file_name_var.get().strip()
        if not file_name:
            self.error_label["text"] = "File name is required for export."
            self.file_name_entry["background"] = "red"
            return
        else:
            self.error_label["text"] = ""
            self.file_name_entry["background"] = "white"

        self.replace_text_in_document()
        folder = filedialog.askdirectory()
        word_export_path = os.path.join(folder, f"{file_name}.docx")
        self.pdf_export_path = os.path.join(folder, f"{file_name}.pdf")
        self.new_doc.save(word_export_path)

        # Wait for the file to be saved
        time_to_wait = 10
        time_counter = 0
        while not os.path.exists(word_export_path) or os.path.getsize(word_export_path) == 0:
            time.sleep(1)
            time_counter += 1
            if time_counter > time_to_wait:
                self.error_label["text"] = "Timed out waiting for file to be saved."
                return

        try:
            # Convert to PDF using pywin32
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False

            doc = word.Documents.Open(word_export_path)
            doc.SaveAs(self.pdf_export_path, FileFormat=17)  # 17 represents the file format for PDF in Word
            doc.Close()
            word.Quit()

            self.error_label["text"] = f"Exported to {self.pdf_export_path}"

        except Exception as e:
            self.error_label["text"] = f"An error occurred: {e}"


    # Preview the exported PDF
    def preview(self):
        subprocess.run(['start', 'cmd', '/c', self.pdf_export_path], shell=True)

if __name__ == "__main__":
    root = tk.Tk()
    app = BookApp(root)
    root.mainloop()